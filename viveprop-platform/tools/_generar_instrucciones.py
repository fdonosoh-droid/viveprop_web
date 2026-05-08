"""
_generar_instrucciones.py — Genera Docs/INSTRUCCIONES_ACTUALIZACIONES.docx
Uso: python tools/_generar_instrucciones.py
"""
import sys
from pathlib import Path
from datetime import date

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent.parent
OUT  = ROOT / "Docs" / "INSTRUCCIONES_ACTUALIZACIONES.docx"

# ── Colores corporativos ──────────────────────────────────────────────────────
AZUL       = RGBColor(0x1B, 0x3A, 0x6B)   # azul oscuro
VERDE      = RGBColor(0x1E, 0x7E, 0x34)   # verde (automático)
NARANJA    = RGBColor(0xC7, 0x57, 0x00)   # naranja (manual)
GRIS       = RGBColor(0x55, 0x55, 0x55)   # gris texto
GRIS_CLARO = RGBColor(0xF5, 0xF5, 0xF5)

# ── Helpers ───────────────────────────────────────────────────────────────────

def shade_cell(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = AZUL
    return p

def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = AZUL
    return p

def add_h3(doc, text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = GRIS
    return p

def add_body(doc, text, bold=False, color=None):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    for run in p.runs:
        run.font.size = Pt(11)
        if bold:  run.bold = True
        if color: run.font.color.rgb = color
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.6)
    return p

def add_numbered(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Number")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.6)
    return p

def add_badge(doc, text, color: RGBColor, label: str):
    """Párrafo con etiqueta de tipo AUTOMÁTICO / MANUAL."""
    p = doc.add_paragraph()
    run = p.add_run(f"  {label}  ")
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size  = Pt(10)
    run.font.highlight_color = None
    # Fondo simulado con sombreado (workaround: usar tabla de 1 celda)
    return p

def add_table_row(table, cells, header=False):
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = text
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(10)
            if header:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        if header:
            shade_cell(cell, "1B3A6B")
    return row

def add_info_box(doc, title: str, lines: list[str], border_color="1E7E34"):
    """Cuadro de información destacada usando tabla de 1 columna."""
    tbl = doc.add_table(rows=0, cols=1)
    tbl.style = "Table Grid"
    # Título
    r0 = tbl.add_row()
    c0 = r0.cells[0]
    c0.text = title
    shade_cell(c0, border_color)
    for run in c0.paragraphs[0].runs:
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
    # Contenido
    r1 = tbl.add_row()
    c1 = r1.cells[0]
    c1.text = ""
    shade_cell(c1, "F5F5F5")
    for line in lines:
        p = c1.add_paragraph(line)
        p.paragraph_format.left_indent = Cm(0.3)
        for run in p.runs:
            run.font.size = Pt(10)
    doc.add_paragraph()  # espacio

# ── Documento ─────────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── Portada ───────────────────────────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run("INSTRUCCIONES DE ACTUALIZACIÓN")
    run.font.size  = Pt(22)
    run.font.bold  = True
    run.font.color.rgb = AZUL

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p_sub.add_run("Plataforma ViveProp — Guía operativa de flujos de datos")
    run2.font.size  = Pt(13)
    run2.font.color.rgb = GRIS

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p_date.add_run(f"Actualizado: {date.today().strftime('%d de %B de %Y')}")
    run3.font.size  = Pt(10)
    run3.font.color.rgb = GRIS

    doc.add_paragraph()

    # ── Introducción ─────────────────────────────────────────────────────────
    add_h1(doc, "1. Visión General")

    add_body(doc,
        "La plataforma ViveProp consume datos de cuatro fuentes distintas. "
        "Cada fuente tiene su propio flujo de actualización — algunos automáticos "
        "y otros manuales. Este documento describe cada flujo, su frecuencia y los "
        "pasos exactos para actualizarlo.")

    doc.add_paragraph()

    # Tabla resumen
    add_h3(doc, "Resumen de flujos")
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    add_table_row(tbl, ["Flujo", "Frecuencia", "Tipo", "Acción requerida"], header=True)
    rows_data = [
        ("Stock Secundario\n(Propiedades Usadas)",  "Diaria — 06:00 AM",        "AUTOMÁTICO",  "Ninguna (GitHub Actions)"),
        ("Stock y Proyectos Primario\n(Mercado Nuevo)", "Cuando hay nueva lista", "MANUAL",      "Doble clic en actualizar_primario.bat"),
        ("Condiciones Comerciales",                  "Cuando hay cambios de CC",  "MANUAL",      "Editar planilla + doble clic en actualizar_cc.bat"),
        ("Fotos y Multimedia\n(Proyectos Primario)", "Al agregar fotos nuevas",   "MANUAL",      "Copiar fotos + correr copy_portadas.py"),
    ]
    for data in rows_data:
        r = tbl.add_row()
        for i, text in enumerate(data):
            cell = r.cells[i]
            cell.text = text
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
                if i == 2:
                    run.bold = True
                    run.font.color.rgb = VERDE if text == "AUTOMÁTICO" else NARANJA
    doc.add_paragraph()

    add_info_box(doc, "REGLA GENERAL — CÓMO SE PUBLICA UN CAMBIO", [
        "Todos los flujos manuales terminan con git commit + git push.",
        "Vercel detecta el push y despliega automáticamente en 1-2 minutos.",
        "Los .bat incluyen el commit y push — no es necesario hacerlo a mano.",
    ], border_color="1B3A6B")

    doc.add_page_break()

    # ── FLUJO 1: Stock Secundario ─────────────────────────────────────────────
    add_h1(doc, "2. Stock Secundario — Propiedades Usadas (AssetPlan)")

    p = doc.add_paragraph()
    r1 = p.add_run("TIPO: ")
    r1.bold = True
    r2 = p.add_run("AUTOMÁTICO — no requiere intervención")
    r2.font.color.rgb = VERDE
    r2.bold = True

    doc.add_paragraph()
    add_h2(doc, "2.1 Qué actualiza")
    add_bullet(doc, "Cartera de propiedades usadas disponibles para venta/arriendo")
    add_bullet(doc, "Precios, tipologías, bono pie, estado de administración")
    add_bullet(doc, "~970 propiedades actualizadas diariamente")

    add_h2(doc, "2.2 Fuente de datos")
    add_info_box(doc, "FUENTE LIVE", [
        "Google Sheets público de AssetPlan (acceso sin autenticación)",
        "Los datos se reflejan en el programa con máximo 24h de desfase.",
    ], border_color="1E7E34")

    add_h2(doc, "2.3 Cómo funciona (automático)")
    add_bullet(doc, "GitHub Actions ejecuta el workflow 'Actualizar Stock Secundario' todos los días a las 06:00 AM hora Chile.")
    add_bullet(doc, "El script descarga el Google Sheet como CSV, mapea los campos y genera stock.json.")
    add_bullet(doc, "Si hay cambios, hace commit y push automático. Vercel despliega en 1-2 min.")
    add_bullet(doc, "Si no hay cambios, el workflow termina sin hacer nada.")

    add_h2(doc, "2.4 Cómo verificar que está funcionando")
    add_numbered(doc, "Ir a GitHub → repositorio viveprop_web → pestaña Actions")
    add_numbered(doc, "Buscar el workflow 'Actualizar Stock Secundario'")
    add_numbered(doc, "Ver que el último run fue exitoso (ícono verde)")

    add_h2(doc, "2.5 Cómo forzar una actualización manual")
    add_info_box(doc, "ACTUALIZACIÓN MANUAL (si se necesita antes de las 06:00 AM)", [
        "Opción A — desde GitHub:",
        "  1. Ir a Actions → Actualizar Stock Secundario → Run workflow",
        "",
        "Opción B — desde la PC local:",
        "  1. Doble clic en:  viveprop-platform\\tasks\\actualizar_stock.bat",
        "  (Requiere Python 3.14 y conexión a internet)",
    ], border_color="1E7E34")

    doc.add_paragraph()
    add_body(doc, "Archivos involucrados:", bold=True)
    add_bullet(doc, "Script:        tools/gsheet_to_stock.py")
    add_bullet(doc, "Workflow:      .github/workflows/actualizar_stock.yml")
    add_bullet(doc, "Bat local:     tasks/actualizar_stock.bat")
    add_bullet(doc, "Output JSON:   frontend/public/data/stock.json")

    doc.add_page_break()

    # ── FLUJO 2: Stock Primario ───────────────────────────────────────────────
    add_h1(doc, "3. Stock y Proyectos Primario — Mercado Nuevo")

    p = doc.add_paragraph()
    r1 = p.add_run("TIPO: ")
    r1.bold = True
    r2 = p.add_run("MANUAL — ejecutar cada vez que las inmobiliarias envíen nuevas listas")
    r2.font.color.rgb = NARANJA
    r2.bold = True

    doc.add_paragraph()
    add_h2(doc, "3.1 Qué actualiza")
    add_bullet(doc, "Disponibilidad de unidades (disponible / reservado / vendido)")
    add_bullet(doc, "Precios de lista de cada unidad")
    add_bullet(doc, "Información de proyectos: entrega, dirección, comuna")
    add_bullet(doc, "~84 proyectos y ~10.000 unidades")

    add_h2(doc, "3.2 Fuente de datos")
    add_info_box(doc, "FUENTE DE DATOS", [
        "Base de datos Access:  H:\\Mi unidad\\Viveprop\\00 - DDBB_AGENCIA\\STOCK MERCADO PRIMARIO\\STOCK MP.accdb",
        "",
        "REQUISITO: la PC debe tener Google Drive montado en H:\\ al momento de correr el bat.",
        "La base Access es actualizada en forma independiente por cada inmobiliaria.",
    ], border_color="C75700")

    add_h2(doc, "3.3 Instrucción de actualización")
    add_info_box(doc, "PASOS — ACTUALIZACIÓN STOCK PRIMARIO", [
        "1. Verificar que Google Drive esté montado (unidad H:\\ visible en el Explorador)",
        "2. Doble clic en:   viveprop-platform\\tasks\\actualizar_primario.bat",
        "3. El bat:",
        "      a. Conecta a la base Access y lee la query STOCK_MP_VIVEPROP_WEB",
        "      b. Genera projects.json con ~10.000 unidades",
        "      c. Hace git commit y git push automáticamente",
        "4. Vercel despliega en 1-2 minutos",
        "5. Verificar en el programa que los cambios se reflejan",
    ], border_color="C75700")

    add_h2(doc, "3.4 Agregar un proyecto nuevo")
    add_body(doc,
        "Cuando se incorpora un proyecto nuevo al Access que aún no existe en el programa:")
    add_numbered(doc, "Confirmar que el proyecto ya está cargado en la base Access con su NEMOTECNICO definido.")
    add_numbered(doc, "Abrir el archivo tools/access_to_projects.py y agregar el NEMOTECNICO al diccionario NEMO_MAP:")
    p_code = doc.add_paragraph('    "NEMO": "proj_nombre-del-proyecto",')
    p_code.paragraph_format.left_indent = Cm(1.5)
    for run in p_code.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(9)
    add_numbered(doc, "Agregar el proyecto en data/projects.xlsx (hoja Proyectos) con su id, nombre, inmobiliaria y descripción.")
    add_numbered(doc, "Correr actualizar_primario.bat.")

    add_h2(doc, "3.5 Archivos involucrados")
    add_bullet(doc, "Script:        tools/access_to_projects.py")
    add_bullet(doc, "Bat:           tasks/actualizar_primario.bat")
    add_bullet(doc, "Datos base:    data/projects.xlsx  (info estática del proyecto)")
    add_bullet(doc, "Fotos base:    frontend/public/data/fotos_pri.json")
    add_bullet(doc, "Output JSON:   frontend/public/data/projects.json")

    doc.add_page_break()

    # ── FLUJO 3: Condiciones Comerciales ──────────────────────────────────────
    add_h1(doc, "4. Condiciones Comerciales (CC)")

    p = doc.add_paragraph()
    r1 = p.add_run("TIPO: ")
    r1.bold = True
    r2 = p.add_run("MANUAL — ejecutar cuando las inmobiliarias informen cambios de CC")
    r2.font.color.rgb = NARANJA
    r2.bold = True

    doc.add_paragraph()
    add_h2(doc, "4.1 Qué actualiza")
    add_bullet(doc, "Descuentos por proyecto y tipología")
    add_bullet(doc, "Aportes y bonos pie de la inmobiliaria")
    add_bullet(doc, "Esquema de cuotas del pie y pie en construcción")
    add_bullet(doc, "Reserva (monto en CLP o en UF)")
    add_bullet(doc, "Crédito directo, cuotón y tipo de entrega")
    add_bullet(doc, "Todos los valores que aparecen en gris en el cotizador")

    add_h2(doc, "4.2 Fuente de datos")
    add_info_box(doc, "ARCHIVO FUENTE — PLANILLA UNIFICADA", [
        "Archivo:  data/condiciones_cc_unified.xlsx",
        "",
        "Una sola planilla con una fila por proyecto.",
        "Columnas principales:",
        "  proyecto_id        — slug del proyecto (ej: proj_vicuna-mackenna-7589-ii)",
        "  inmobiliaria       — MAESTRA / INGEVEC / RVC / TOCTOC / URMENETA",
        "  descuento_depto    — % descuento sobre el depto (ej: 10 = 10%)",
        "  aporte_inmobiliaria— % bono / certificado de pago de la inmobiliaria",
        "  reserva_clp        — monto reserva en pesos (ej: 100000)",
        "  reserva_uf         — monto reserva en UF (si aplica; reserva_clp se calcula automático)",
        "  cuotas_pie_n       — cantidad de cuotas del saldo del pie",
        "  pie_pct_default    — % pie total sugerido (vacío = usa default del sistema)",
        "  pie_const_pct      — % pie período construcción",
        "  credito_directo_pct— % crédito directo de la inmobiliaria",
        "  cuoton_pct         — % cuotón",
        "  tipo_entrega       — texto libre (ej: 2do semestre 2027)",
        "  nota               — texto adicional visible en el cotizador",
    ], border_color="C75700")

    add_h2(doc, "4.3 Instrucción de actualización")
    add_info_box(doc, "PASOS — ACTUALIZACIÓN CONDICIONES COMERCIALES", [
        "1. Abrir:   data/condiciones_cc_unified.xlsx",
        "2. Modificar los valores que cambiaron (una fila = un proyecto)",
        "   - Los porcentajes van como números: 10 para 10%, no '10%'",
        "   - Las reservas en CLP van como enteros: 100000 (sin puntos de miles)",
        "   - Si la reserva es en UF (ej: TOCTOC 10 UF): poner 10 en reserva_uf y dejar vacío reserva_clp",
        "     El programa calculará el CLP automáticamente con la UF del día al correr el bat",
        "3. Guardar y cerrar el Excel",
        "4. Doble clic en:   viveprop-platform\\tasks\\actualizar_cc.bat",
        "5. El bat:",
        "      a. Corre cc_importer.py (obtiene UF de mindicador.cl si hay reservas en UF)",
        "      b. Genera cc.json con valores normalizados",
        "      c. Hace git commit y git push automáticamente",
        "6. Vercel despliega en 1-2 minutos",
    ], border_color="C75700")

    add_h2(doc, "4.4 Agregar un proyecto nuevo a las CC")
    add_numbered(doc, "Abrir data/condiciones_cc_unified.xlsx")
    add_numbered(doc, "Agregar una fila nueva al final con los datos del proyecto")
    add_body(doc, "     El campo proyecto_id debe contener el slug exacto (ej: proj_nombre-proyecto).")
    add_body(doc, "     El slug está en frontend/public/data/projects.json → campo 'id'.")
    add_numbered(doc, "Guardar y correr actualizar_cc.bat")

    add_h2(doc, "4.5 Agregar una inmobiliaria nueva")
    add_body(doc,
        "No requiere cambios de código. El campo inmobiliaria en la planilla es texto libre. "
        "Simplemente usar el nombre de la nueva inmobiliaria en la columna correspondiente "
        "y agregar los proyectos con sus datos normalizados.")

    add_h2(doc, "4.6 Archivos involucrados")
    add_bullet(doc, "Planilla fuente:   data/condiciones_cc_unified.xlsx")
    add_bullet(doc, "Script:            tools/cc_importer.py")
    add_bullet(doc, "Bat:               tasks/actualizar_cc.bat")
    add_bullet(doc, "Output JSON:       frontend/public/data/cc.json")

    doc.add_page_break()

    # ── FLUJO 4: Fotos y Multimedia ───────────────────────────────────────────
    add_h1(doc, "5. Fotos y Multimedia — Proyectos Primario")

    p = doc.add_paragraph()
    r1 = p.add_run("TIPO: ")
    r1.bold = True
    r2 = p.add_run("MANUAL — ejecutar al agregar fotos, tipologías o PDFs nuevos")
    r2.font.color.rgb = NARANJA
    r2.bold = True

    doc.add_paragraph()
    add_h2(doc, "5.1 Qué actualiza")
    add_bullet(doc, "Foto de portada de cada proyecto")
    add_bullet(doc, "Galería de imágenes (fotos internas, renders, etc.)")
    add_bullet(doc, "Tipologías (imágenes de plantas de cada tipo de departamento)")
    add_bullet(doc, "PDFs (brochures, fichas técnicas)")

    add_h2(doc, "5.2 Estructura de carpetas")
    add_info_box(doc, "ESTRUCTURA DE FOTOS", [
        "photos/primario/",
        "  └── proj_nombre-del-proyecto/",
        "        ├── portada.jpg          ← foto principal del proyecto",
        "        ├── galeria/             ← imágenes adicionales",
        "        │     ├── foto1.jpg",
        "        │     └── foto2.jpg",
        "        ├── tipologias/          ← plantas por tipología",
        "        │     ├── 1d1b.jpg",
        "        │     └── 2d1b.jpg",
        "        └── pdfs/               ← brochures y fichas",
        "              └── brochure.pdf",
    ], border_color="1B3A6B")

    add_h2(doc, "5.3 Instrucción — agregar fotos a un proyecto")
    add_info_box(doc, "PASOS — ACTUALIZACIÓN DE FOTOS", [
        "1. Copiar las fotos a la carpeta correspondiente en photos/primario/proj_xxx/",
        "   (crear subcarpetas galeria/, tipologias/, pdfs/ si no existen)",
        "2. Desde el directorio viveprop-platform/, correr:",
        "      C:\\Python314\\python.exe tools\\copy_portadas.py",
        "3. El script:",
        "      a. Copia las fotos a frontend/public/photos/pri/",
        "      b. Actualiza el manifest fotos_pri.json",
        "4. Para actualizar rutas de PDFs:",
        "      C:\\Python314\\python.exe tools\\update_pdf_paths.py",
        "5. Hacer git add + commit + push manualmente:",
        "      git add frontend/public/photos/pri/ frontend/public/data/fotos_pri.json",
        "      git commit -m 'chore: actualiza fotos proyecto X'",
        "      git push",
    ], border_color="C75700")

    add_h2(doc, "5.4 Archivos involucrados")
    add_bullet(doc, "Fuente:         photos/primario/proj_xxx/")
    add_bullet(doc, "Script fotos:   tools/copy_portadas.py")
    add_bullet(doc, "Script PDFs:    tools/update_pdf_paths.py")
    add_bullet(doc, "Manifest:       frontend/public/data/fotos_pri.json")
    add_bullet(doc, "Output:         frontend/public/photos/pri/proj_xxx/")

    doc.add_page_break()

    # ── ACTUALIZACIÓN MASIVA ──────────────────────────────────────────────────
    add_h1(doc, "6. Actualización Masiva")

    add_body(doc,
        "Cuando se necesita actualizar todo el programa de una vez (ej: inicio de mes, "
        "después de vacaciones, o al recibir múltiples actualizaciones simultáneas), "
        "seguir el orden recomendado:")

    doc.add_paragraph()

    add_info_box(doc, "ORDEN RECOMENDADO — ACTUALIZACIÓN MASIVA", [
        "PASO 1 — Stock Primario (si hay nueva lista de precios o disponibilidad)",
        "  → Doble clic en:   tasks\\actualizar_primario.bat",
        "  → Esperar que termine (10-30 segundos)",
        "",
        "PASO 2 — Condiciones Comerciales (si hay cambios de CC)",
        "  → Editar:          data\\condiciones_cc_unified.xlsx",
        "  → Doble clic en:   tasks\\actualizar_cc.bat",
        "  → Esperar que termine (5-15 segundos)",
        "",
        "PASO 3 — Fotos y Multimedia (si hay fotos nuevas)",
        "  → Copiar fotos a   photos\\primario\\proj_xxx\\",
        "  → Correr:          python tools\\copy_portadas.py",
        "  → Hacer git push manualmente",
        "",
        "PASO 4 — Stock Secundario",
        "  → No requiere acción — se actualiza solo a las 06:00 AM",
        "  → Si se necesita actualizar ahora: ir a GitHub → Actions → Run workflow",
        "",
        "NOTA: El programa en Vercel refleja cada cambio en 1-2 minutos después del push.",
    ], border_color="1B3A6B")

    doc.add_paragraph()

    add_h2(doc, "6.1 Tiempos de propagación")
    tbl2 = doc.add_table(rows=1, cols=3)
    tbl2.style = "Table Grid"
    add_table_row(tbl2, ["Acción", "Tiempo hasta visible en el programa", "Observación"], header=True)
    times = [
        ("Correr cualquier bat",       "1-2 minutos",       "Vercel despliega al detectar el push"),
        ("GitHub Actions (stock)",     "1-3 minutos",       "Incluye tiempo de ejecución del workflow"),
        ("Editar solo la planilla CC", "No se publica",     "Siempre correr el bat después de editar"),
        ("Agregar fotos",              "1-2 min tras push", "Recordar hacer git push manual"),
    ]
    for data in times:
        r = tbl2.add_row()
        for i, text in enumerate(data):
            cell = r.cells[i]
            cell.text = text
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_page_break()

    # ── RESOLUCIÓN DE PROBLEMAS ───────────────────────────────────────────────
    add_h1(doc, "7. Resolución de Problemas Frecuentes")

    problemas = [
        (
            "El bat de primario falla con error de conexión a Access",
            [
                "Verificar que Google Drive esté montado como unidad H:\\",
                "Abrir el Explorador de Windows → ver si aparece la unidad H:\\",
                "Si no aparece: abrir Google Drive app y esperar sincronización",
                "Si el archivo .accdb está abierto en Access: cerrarlo antes de correr el bat",
            ]
        ),
        (
            "El bat de CC termina sin hacer push ('Sin cambios en CC')",
            [
                "Esto es normal si se corrió el bat sin haber modificado la planilla",
                "Si se hicieron cambios y no se reflejan: verificar que se guardó el Excel antes de correr el bat",
            ]
        ),
        (
            "Las condiciones comerciales no aparecen en el cotizador (todos los campos en 0%)",
            [
                "Verificar que el proyecto_id en condiciones_cc_unified.xlsx coincide exactamente con el id en projects.json",
                "El id es sensible a mayúsculas y debe tener el prefijo proj_",
                "Correr actualizar_cc.bat y verificar que el proyecto aparece en la salida",
            ]
        ),
        (
            "GitHub Actions falla (ícono rojo en pestaña Actions)",
            [
                "Ir a Actions → click en el run fallido → ver el log de error",
                "Causa más frecuente: cambio en el formato del Google Sheet de AssetPlan",
                "En ese caso: revisar el script gsheet_to_stock.py y actualizar los nombres de columna",
            ]
        ),
        (
            "Las fotos de un proyecto no aparecen en el programa",
            [
                "Verificar que copy_portadas.py se corrió después de copiar las fotos",
                "Verificar que se hizo git push después de correr el script",
                "El nombre de la carpeta debe coincidir exactamente con el id del proyecto (proj_xxx)",
            ]
        ),
    ]

    for titulo, pasos in problemas:
        add_h3(doc, f"✦  {titulo}")
        for paso in pasos:
            add_bullet(doc, paso)
        doc.add_paragraph()

    # ── Nota final ────────────────────────────────────────────────────────────
    doc.add_page_break()
    add_h1(doc, "8. Referencia Rápida — Comandos")

    add_info_box(doc, "COMANDOS ÚTILES (ejecutar desde viveprop-platform/)", [
        "# Actualizar stock primario:",
        "  tasks\\actualizar_primario.bat",
        "",
        "# Actualizar condiciones comerciales:",
        "  tasks\\actualizar_cc.bat",
        "",
        "# Actualizar fotos:",
        "  C:\\Python314\\python.exe tools\\copy_portadas.py",
        "  C:\\Python314\\python.exe tools\\update_pdf_paths.py",
        "",
        "# Forzar actualización de stock secundario:",
        "  tasks\\actualizar_stock.bat",
        "",
        "# Ver qué archivos cambiaron antes de hacer push:",
        "  git status",
        "  git diff --stat",
    ], border_color="1B3A6B")

    doc.save(OUT)
    print(f"[ok] {OUT}")


if __name__ == "__main__":
    build()
