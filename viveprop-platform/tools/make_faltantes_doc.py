# -*- coding: utf-8 -*-
"""Genera PROYECTOS_FALTANTES.docx en la carpeta Docs/"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT = Path(__file__).parent.parent / "Docs" / "PROYECTOS_FALTANTES.docx"

doc = Document()

# ── Márgenes más estrechos ────────────────────────────────────
for section in doc.sections:
    section.left_margin   = int(1.5 * 914400 / 2.54)
    section.right_margin  = int(1.5 * 914400 / 2.54)
    section.top_margin    = int(2.0 * 914400 / 2.54)
    section.bottom_margin = int(2.0 * 914400 / 2.54)

# ── Helpers ───────────────────────────────────────────────────

def shading(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_table(doc, headers, rows, hdr_fill="1A56AA", alt_fill="EEF3FB"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    for ci, txt in enumerate(headers):
        c = table.rows[0].cells[ci]
        c.text = txt
        run = c.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading(c, hdr_fill)

    # Data rows
    for ri, row_data in enumerate(rows):
        fill = alt_fill if ri % 2 == 0 else "FFFFFF"
        for ci, txt in enumerate(row_data):
            c = table.rows[ri + 1].cells[ci]
            c.text = txt
            c.paragraphs[0].runs[0].font.size = Pt(9)
            shading(c, fill)

    doc.add_paragraph("")


def add_hr(doc):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "BBBBBB")
    bdr.append(bot)
    pPr.append(bdr)


def heading1(doc, text, hex_color):
    h = doc.add_heading(text, level=1)
    r = h.runs[0]
    r.font.color.rgb = RGBColor(*bytes.fromhex(hex_color))
    r.font.size = Pt(13)


def heading2(doc, text, hex_color):
    h = doc.add_heading(text, level=2)
    r = h.runs[0]
    r.font.color.rgb = RGBColor(*bytes.fromhex(hex_color))
    r.font.size = Pt(11)


# ═══════════════════════════════════════════════════════════════
# TÍTULO
# ═══════════════════════════════════════════════════════════════
t = doc.add_heading("Proyectos Faltantes de Incorporar", level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.runs[0].font.color.rgb = RGBColor(0x1A, 0x56, 0xAA)
t.runs[0].font.size = Pt(18)

sub = doc.add_paragraph("Fotos y datos de portada pendientes de agregar a photos/primario/")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(10)
sub.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
sub.runs[0].italic = True
doc.add_paragraph("")


# ═══════════════════════════════════════════════════════════════
# GRUPO A — Carpeta existe, falta contenido o portada (9 proyectos)
# ═══════════════════════════════════════════════════════════════
heading2(doc,
    "Grupo A — Carpeta ya existe, falta portada o fotos  (9 proyectos)",
    "D97300")

note = doc.add_paragraph(
    "Estas carpetas ya están en photos/primario/. "
    "Agregar la imagen de portada, asignar la ruta en la columna foto_portada del Excel "
    "y re-ejecutar las herramientas."
)
note.runs[0].font.size = Pt(9)
note.runs[0].italic = True
doc.add_paragraph("")

# INGEVEC
heading1(doc, "INGEVEC", "215D9A")
add_table(doc,
    ["Proyecto", "Ruta de carpeta  (relativa a viveprop-platform/photos/primario/)", "Estado"],
    [
        ["SANTOS OSSA",
         "INGEVEC(13)/SANTOS OSSA (STO)/",
         "Tiene fotos — falta foto_portada en Excel"],
        ["NUEVA ESMERALDA",
         "INGEVEC(13)/NUEVA ESMERALDA (NES)/",
         "Tiene fotos — falta foto_portada en Excel"],
        ["VESPUCIO CAPITAL",
         "INGEVEC(13)/VESPUCIO CAPITAL (VES)/",
         "Tiene fotos — falta foto_portada en Excel"],
    ],
)

# RVC
heading1(doc, "RVC", "215D9A")
add_table(doc,
    ["Proyecto", "Ruta de carpeta  (relativa a viveprop-platform/photos/primario/)", "Estado"],
    [
        ["CONDOMINIO FRANCISCO ZELADA TORRE B",
         "RVC(4)/FRANCISCO ZELADA TORRE B (FZB)/",
         "Tiene fotos hash — falta foto_portada en Excel"],
        ["EDIFICIO PARQUE URBANO",
         "RVC(4)/PARQUE URBANO VINA (PUV)/",
         "Carpeta vacía — agregar fotos"],
    ],
)

# URMENETA
heading1(doc, "URMENETA", "215D9A")
add_table(doc,
    ["Proyecto", "Ruta de carpeta  (relativa a viveprop-platform/photos/primario/)", "Estado"],
    [
        ["Barrio Cueto",
         "URMENETA(20)/Barrio Cueto(3)-ENTREGA FUTURA(BONO PIE 10%)-CUETO 1218,SANTIAGO/",
         "Solo PDF/xlsx — agregar fotos"],
        ["Ciudad Siete",
         "URMENETA(20)/Ciudad Siete(3)-ENTREGA FUTURA(GASTO COMUN GRATIS POR UN ANO)-AV SIETE 2965, CERRILLOS/",
         "Solo PDF — agregar fotos"],
        ["Fuenzalida Urrejola 303",
         "URMENETA(20)/FUENZALIDA URREJOLA-ENTREGA FUTURA(Fuenzalida Urrejola 303, La Cisterna)/",
         "Subcarpetas vacias — agregar fotos"],
        ["Garden Macul",
         "URMENETA(20)/Garden Macul(3)-ENTREGA INMEDIATA-MACUL 2320, MACUL/",
         "Carpeta vacia — agregar fotos"],
    ],
)

add_hr(doc)
doc.add_paragraph("")


# ═══════════════════════════════════════════════════════════════
# GRUPO B — Sin carpeta, hay que crearla (13 proyectos)
# ═══════════════════════════════════════════════════════════════
heading2(doc,
    "Grupo B — Sin carpeta, debe crearse  (13 proyectos)",
    "C02A2A")

note2 = doc.add_paragraph(
    "Estas carpetas no existen. Crear la carpeta con el nombre sugerido dentro de "
    "photos/primario/{INMOBILIARIA}/, agregar portada y fotos, luego asignar foto_portada en Excel."
)
note2.runs[0].font.size = Pt(9)
note2.runs[0].italic = True
doc.add_paragraph("")

# INGEVEC
heading1(doc, "INGEVEC", "C02A2A")
add_table(doc,
    ["Proyecto", "Nombre de carpeta sugerido", "Accion"],
    [
        ["ZANARTU 1111", "INGEVEC(13)/ZANARTU 1111 (ZAN)/", "Crear carpeta + agregar fotos"],
        ["SANTA ISABEL",  "INGEVEC(13)/SANTA ISABEL (SIS)/",  "Crear carpeta + agregar fotos"],
        ["SANTA ROSA",    "INGEVEC(13)/SANTA ROSA (SRO)/",    "Crear carpeta + agregar fotos"],
        ["CURAUMA",       "INGEVEC(13)/CURAUMA (CUR)/",       "Crear carpeta + agregar fotos"],
    ],
    hdr_fill="8B1A1A",
)

# MAESTRA
heading1(doc, "MAESTRA INMOBILIARIA", "C02A2A")
add_table(doc,
    ["Proyecto", "Nombre de carpeta sugerido", "Accion"],
    [
        ["DISTRITO CENTRO",
         "MAESTRA INMOBILIARIA(15)/Distrito Centro/",
         "Crear carpeta + agregar fotos"],
    ],
    hdr_fill="8B1A1A",
)

# URMENETA
heading1(doc, "URMENETA", "C02A2A")
add_table(doc,
    ["Proyecto", "Nombre de carpeta sugerido", "Accion"],
    [
        ["JMC 608",             "URMENETA(20)/JMC 608/",             "Crear carpeta + agregar fotos"],
        ["Mapocho 2880",        "URMENETA(20)/Mapocho 2880/",        "Crear carpeta + agregar fotos"],
        ["NeoFlorida 3",        "URMENETA(20)/NeoFlorida 3/",        "Crear carpeta + agregar fotos"],
        ["Status",              "URMENETA(20)/Status/",              "Crear carpeta + agregar fotos"],
        ["Suena Lincoyan",      "URMENETA(20)/Suena Lincoyan/",      "Crear carpeta + agregar fotos"],
        ["Suena Marin",         "URMENETA(20)/Suena Marin/",         "Crear carpeta + agregar fotos"],
        ["Suena Santa Elisa",   "URMENETA(20)/Suena Santa Elisa/",   "Crear carpeta + agregar fotos"],
        ["Suena Toesca",        "URMENETA(20)/Suena Toesca/",        "Crear carpeta + agregar fotos"],
    ],
    hdr_fill="8B1A1A",
)

add_hr(doc)
doc.add_paragraph("")


# ═══════════════════════════════════════════════════════════════
# FLUJO DE TRABAJO
# ═══════════════════════════════════════════════════════════════
heading2(doc, "Flujo para activar proyectos", "1A56AA")

steps = [
    ("1.", "Colocar portada (y galería/tipologías opcionales) dentro de la carpeta indicada."),
    ("2.", "Abrir el Excel y apuntar la columna foto_portada a la ruta de la portada."),
    ("",   "    Ej: photos/primario/INGEVEC(13)/ZANARTU 1111 (ZAN)/portada.jpg"),
    ("3.", "Ejecutar:  python tools/copy_portadas.py"),
    ("4.", "Ejecutar:  python tools/excel_to_js.py"),
    ("5.", "git add → git commit → git push  (Vercel despliega automáticamente)."),
]
for num, text in steps:
    p = doc.add_paragraph()
    if num:
        r = p.add_run(num + "  ")
        r.bold = True
        r.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    if text.startswith("    "):
        r2.font.color.rgb = RGBColor(0x44, 0x44, 0xAA)
        r2.italic = True

# ── Guardar ───────────────────────────────────────────────────
doc.save(OUT)
print(f"Guardado: {OUT}")
