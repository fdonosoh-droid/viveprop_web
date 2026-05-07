# ANTES DE EJECUTAR EJECUTAR EL POWERSHEL EN TERMINAL
# python -m pip install requests
# python -m pip install requests pandas tqdm python-docx openpyxl


import os
import json
import requests
import pandas as pd
from tqdm import tqdm
from pathlib import Path
from docx import Document
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed

# ================= CONFIGURACION GENERAL =================
EXCEL_PATH = "cartera_disponible_assetplan.xlsx"
BASE_NAME = "propiedades_assetplan"
API_URL = "https://api.assetplan.cl/api/v1/property/for_sale/{id}?list=1"
MAX_WORKERS = 10
BATCH_SIZE = 20
CONTROL_FILE = "control_versiones.xlsx"
RESUMEN_CAMBIOS = "resumen_cambios_versiones.xlsx"

CAMPOS_CLAVE = [
    "valor_uf",
    "monto_estimado",
    "for_sale",
    "fecha_disponible_venta",
    "retorno",
    "arriendo"
]

# ================= FUNCIONES AUXILIARES =================
def limpiar_nombre(texto):
    if not texto:
        return ""
    return (
        str(texto)
        .replace("/", "-")
        .replace("\\", "-")
        .replace(":", "")
        .replace("*", "")
        .replace("?", "")
        .replace('"', "")
        .replace("<", "")
        .replace(">", "")
        .replace("|", "")
        .strip()
    )

def format_value(value):
    if isinstance(value, dict):
        return ", ".join([f"{k}: {v}" for k, v in value.items()])
    if isinstance(value, list):
        return ", ".join([str(v) for v in value])
    return "" if value is None else str(value)

def flatten_json(data, prefix=""):
    out = {}
    if isinstance(data, dict):
        for k, v in data.items():
            out.update(flatten_json(v, f"{prefix}{k}_"))
    elif isinstance(data, list):
        for i, v in enumerate(data):
            out.update(flatten_json(v, f"{prefix}{i}_"))
    else:
        out[prefix[:-1]] = data
    return out

def descargar_imagen(url, carpeta, session):
    nombre = url.split("/")[-1].split("?")[0]
    destino = os.path.join(carpeta, nombre)
    if os.path.exists(destino):
        return 0
    try:
        r = session.get(url, timeout=15)
        if r.status_code == 200:
            with open(destino, "wb") as f:
                f.write(r.content)
            return 1
    except:
        pass
    return 0

def generar_word(data, carpeta, nombre_archivo):
    ruta = os.path.join(carpeta, f"{nombre_archivo}.docx")
    doc = Document()
    doc.add_heading("Ficha Propiedad AssetPlan", level=1)
    for key, value in data.items():
        doc.add_paragraph(f"{key}: {format_value(value)}")
    doc.save(ruta)

# ================= PROCESAR PROPIEDAD =================
def procesar_propiedad(prop_id, codigo, version_dir, session):
    try:
        url = API_URL.format(id=prop_id)
        resp = session.get(url, timeout=30)
        if resp.status_code != 200:
            return None

        api_data = resp.json()
        if "data" not in api_data or not api_data["data"]:
            return None

        data = api_data["data"][0]

        # Compactar GGCC
        if isinstance(data.get("unitggcc"), dict):
            monto = data["unitggcc"].get("monto")
            if monto is not None:
                data["ggcc"] = monto
            del data["unitggcc"]

        direccion = limpiar_nombre(data.get("direccion"))
        unidad = limpiar_nombre(data.get("unidad"))
        ciudad = limpiar_nombre(data.get("ciudad"))
        codigo = limpiar_nombre(codigo)

        carpeta_nombre = f"{codigo} {direccion} dpto. {unidad} {ciudad}".strip()
        carpeta_path = os.path.join(version_dir, carpeta_nombre)
        Path(carpeta_path).mkdir(parents=True, exist_ok=True)

        # Descargar solo fotos_originales
        fotos = data.get("fotos_originales", [])
        urls = [u for u in fotos if isinstance(u, str) and u.startswith("http")]
        descargadas = sum(descargar_imagen(u, carpeta_path, session) for u in urls)

        with open(os.path.join(carpeta_path, "urls.json"), "w", encoding="utf-8") as f:
            json.dump(urls, f, indent=2, ensure_ascii=False)

        generar_word(data, carpeta_path, carpeta_nombre)

        plano = flatten_json(data)
        plano["imagenes_descargadas"] = descargadas
        plano["carpeta"] = carpeta_nombre
        plano["codigo"] = codigo
        plano["id_assetplan"] = prop_id

        return plano
    except:
        return None

# ================= INICIO PROCESO =================
now = datetime.now()
timestamp = now.strftime("%Y%m%d_%H%M")
VERSION_DIR = f"{BASE_NAME}_{timestamp}"
Path(VERSION_DIR).mkdir(parents=True, exist_ok=True)

df = pd.read_excel(EXCEL_PATH)
df.columns = [c.strip().lower() for c in df.columns]

if not {"id_assetplan", "codigo"}.issubset(df.columns):
    raise ValueError("El Excel debe contener las columnas 'id_assetplan' y 'Codigo'")

mapa = (
    df[["id_assetplan", "codigo"]]
    .dropna()
    .astype(str)
    .set_index("id_assetplan")["codigo"]
    .to_dict()
)

ids = list(mapa.keys())
print(f"\n?? {len(ids)} propiedades cargadas desde {EXCEL_PATH}\n")

# ================= DESCARGA MULTIHILO =================
resultados = []
session = requests.Session()

for i in range(0, len(ids), BATCH_SIZE):
    bloque = ids[i:i+BATCH_SIZE]
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        futures = [
            executor.submit(procesar_propiedad, int(pid), mapa[pid], VERSION_DIR, session)
            for pid in bloque
        ]
        for future in as_completed(futures):
            r = future.result()
            if r:
                resultados.append(r)

session.close()

# ================= CONSOLIDADO =================
df_actual = pd.DataFrame(resultados)
consolidado_path = os.path.join(VERSION_DIR, "consolidado_propiedades.xlsx")
df_actual.to_excel(consolidado_path, index=False)

# ================= CONTROL DE VERSIONES =================
if os.path.exists(CONTROL_FILE):
    df_control = pd.read_excel(CONTROL_FILE)
    df_control[f"version_{timestamp}"] = df_actual["id_assetplan"].astype(str)
else:
    df_control = pd.DataFrame({f"version_{timestamp}": df_actual["id_assetplan"].astype(str)})

df_control.to_excel(CONTROL_FILE, index=False)

print("\n? Proceso completado correctamente")
print(f"?? Carpeta version: {VERSION_DIR}")
print(f"?? Consolidado: {consolidado_path}")
print(f"?? Control historico: {CONTROL_FILE}")
